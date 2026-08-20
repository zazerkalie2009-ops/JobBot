import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def build_resume_docx(output_path: str, profile_data: dict, customized_data: dict = None) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        
    primary_color = RGBColor(30, 41, 59)     # #1E293B
    accent_color = RGBColor(37, 99, 235)     # #2563EB
    muted_color = RGBColor(100, 116, 139)    # #64748B
    dark_text = RGBColor(51, 65, 85)         # #334155
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = dark_text
    
    personal = profile_data['personal']
    role = (customized_data and customized_data.get('role')) or profile_data['roles']['default']
    
    # Header: Name
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(1)
    run_name = p_name.add_run(personal['name'])
    run_name.font.size = Pt(19)
    run_name.font.bold = True
    run_name.font.color.rgb = primary_color
    
    # Header: Role
    p_role = doc.add_paragraph()
    p_role.paragraph_format.space_before = Pt(0)
    p_role.paragraph_format.space_after = Pt(3)
    run_role = p_role.add_run(role)
    run_role.font.size = Pt(11.5)
    run_role.font.bold = True
    run_role.font.color.rgb = accent_color
    
    # Header: Contacts
    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(8)
    contacts_str = f"Телефон: {personal['phone']} | Telegram: {personal['telegram']} | Email: {personal['email']} | {personal['location']}"
    run_contact = p_contact.add_run(contacts_str)
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = muted_color
    
    # Summary Box
    summary_text = (customized_data and customized_data.get('summary')) or profile_data['summary_templates']['default']
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.3)
    cell = table.cell(0, 0)
    set_cell_background(cell, 'EFF6FF')
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="2563EB"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p_sum = cell.paragraphs[0]
    p_sum.paragraph_format.space_before = Pt(3)
    p_sum.paragraph_format.space_after = Pt(3)
    p_sum.paragraph_format.left_indent = Inches(0.08)
    p_sum.paragraph_format.right_indent = Inches(0.08)
    
    run_sum = p_sum.add_run(summary_text)
    run_sum.font.size = Pt(9.5)
    run_sum.font.color.rgb = RGBColor(30, 58, 138)
    
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title.upper())
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = primary_color
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)

    # Experience
    add_section_header("Профессиональный опыт")
    for exp in profile_data['experience']:
        p_jh = doc.add_paragraph()
        p_jh.paragraph_format.space_before = Pt(5)
        p_jh.paragraph_format.space_after = Pt(1)
        p_jh.paragraph_format.keep_with_next = True
        
        r_comp = p_jh.add_run(exp['company'] + ' ')
        r_comp.bold = True
        r_comp.font.color.rgb = primary_color
        r_comp.font.size = Pt(10.5)
        
        r_role = p_jh.add_run('| ' + exp['role'] + ' ')
        r_role.bold = True
        r_role.font.color.rgb = accent_color
        r_role.font.size = Pt(10)
        
        r_dates = p_jh.add_run('(' + exp['period'] + ')')
        r_dates.font.color.rgb = muted_color
        r_dates.font.size = Pt(9.5)
        
        if exp.get('description'):
            p_jd = doc.add_paragraph()
            p_jd.paragraph_format.space_before = Pt(0)
            p_jd.paragraph_format.space_after = Pt(2.5)
            p_jd.paragraph_format.keep_with_next = True
            r_desc = p_jd.add_run(exp['description'])
            r_desc.italic = True
            r_desc.font.size = Pt(9)
            r_desc.font.color.rgb = muted_color
            
        for ach in exp['achievements']:
            p_ach = doc.add_paragraph(style='List Bullet')
            p_ach.paragraph_format.space_before = Pt(0)
            p_ach.paragraph_format.space_after = Pt(2)
            
            r_title = p_ach.add_run(ach['title'] + ': ')
            r_title.bold = True
            r_title.font.size = Pt(9.5)
            r_title.font.color.rgb = dark_text
            
            r_text = p_ach.add_run(ach['text'])
            r_text.font.size = Pt(9.5)
            r_text.font.color.rgb = dark_text
                
    # Education
    add_section_header("Образование")
    for edu in profile_data['education']:
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(1.5)
        p_edu.paragraph_format.space_after = Pt(1.5)
        
        r_inst = p_edu.add_run(edu['institution'] + ' ')
        r_inst.bold = True
        r_inst.font.color.rgb = primary_color
        r_inst.font.size = Pt(9.5)
        
        r_deg = p_edu.add_run('— ' + edu['degree'] + ' ')
        r_deg.font.size = Pt(9.5)
        
        r_yr = p_edu.add_run('(' + edu['years'] + ')')
        r_yr.font.color.rgb = muted_color
        r_yr.font.size = Pt(9)
        
    # Skills
    add_section_header("Навыки и стек")
    for sk in profile_data['skills']:
        p_sk = doc.add_paragraph()
        p_sk.paragraph_format.space_before = Pt(1)
        p_sk.paragraph_format.space_after = Pt(1.5)
        
        r_cat = p_sk.add_run(sk['category'] + ': ')
        r_cat.bold = True
        r_cat.font.color.rgb = primary_color
        r_cat.font.size = Pt(9.5)
        
        r_val = p_sk.add_run(', '.join(sk['items']))
        r_val.font.size = Pt(9.5)
        
    # Languages
    p_lang = doc.add_paragraph()
    p_lang.paragraph_format.space_before = Pt(1)
    p_lang.paragraph_format.space_after = Pt(1.5)
    r_l_title = p_lang.add_run("Иностранные языки: ")
    r_l_title.bold = True
    r_l_title.font.color.rgb = primary_color
    r_l_title.font.size = Pt(9.5)
    
    lang_str = ', '.join([f"{l['name']} — {l['level']}" for l in profile_data['languages']])
    r_l_val = p_lang.add_run(lang_str)
    r_l_val.font.size = Pt(9.5)
    
    doc.save(output_path)
    return output_path
